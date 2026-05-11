from pathlib import Path
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm

BASE = Path(__file__).resolve().parent

LINES = [
    "Compétition: Ligue 1 Tunisie",
    "Saison: 2025-2026",
    "Journée: 12",
    "Date: 15/02/2026",
    "Heure: 14h30",
    "Stade: Stade Olympique de Radès",
    "Ville: Radès",
    "Équipe domicile: Club Africain",
    "Équipe extérieure: ES Tunis",
    "Score final: 2-1",
    "Score mi-temps: 1-0",
    "Arbitre central: Mohamed Ben Ali",
    "Assistant 1: Sami Trabelsi",
    "Assistant 2: Hatem Gharsalli",
    "4e arbitre: Anis Zribi",
    "Commissaire: Nabil Zorgati",
    "Entraîneur domicile: Faouzi Benzarti",
    "Entraîneur extérieur: Maher Kanzari",
    "Observations: Match joué devant public. Réserve technique signalée à la 78e minute.",
]

PLAYERS_HEADER = ["Equipe", "Numero", "Joueur", "Licence", "Date naissance", "Poste", "Nationalite", "Statut", "Capitaine", "Gardien"]
PLAYERS = [
    ["Club Africain", 1, "Adam Dhaouadi", "CA-001", "1997-03-11", "Gardien", "TUN", "Titulaire", "", "Oui"],
    ["Club Africain", 2, "Yassine Haddad", "CA-002", "1998-07-19", "Défenseur", "TUN", "Titulaire", "", ""],
    ["Club Africain", 4, "Karim Ben Salem", "CA-004", "1995-01-23", "Défenseur", "TUN", "Titulaire", "", ""],
    ["Club Africain", 5, "Firas Jouini", "CA-005", "1999-04-02", "Défenseur", "TUN", "Titulaire", "", ""],
    ["Club Africain", 8, "Aymen Trabelsi", "CA-008", "1996-09-15", "Milieu", "TUN", "Titulaire", "Oui", ""],
    ["Club Africain", 10, "Ali Mansour", "CA-010", "2000-12-05", "Milieu", "TUN", "Titulaire", "", ""],
    ["Club Africain", 11, "Mohamed Salem", "CA-011", "1998-10-30", "Attaquant", "TUN", "Titulaire", "", ""],
    ["Club Africain", 16, "Soufiane Jaziri", "CA-016", "2001-02-14", "Attaquant", "TUN", "Remplaçant", "", ""],
    ["Club Africain", 17, "Hamza Maaloul", "CA-017", "2002-06-21", "Milieu", "TUN", "Remplaçant", "", ""],
    ["ES Tunis", 1, "Rami Ghanmi", "EST-001", "1994-08-08", "Gardien", "TUN", "Titulaire", "", "Oui"],
    ["ES Tunis", 3, "Oussama Rekik", "EST-003", "1998-01-28", "Défenseur", "TUN", "Titulaire", "", ""],
    ["ES Tunis", 6, "Mahmoud Amri", "EST-006", "1997-05-17", "Défenseur", "TUN", "Titulaire", "", ""],
    ["ES Tunis", 7, "Youssef Gharbi", "EST-007", "2000-11-09", "Milieu", "TUN", "Titulaire", "Oui", ""],
    ["ES Tunis", 9, "Karim Chaabane", "EST-009", "1996-03-25", "Attaquant", "TUN", "Titulaire", "", ""],
    ["ES Tunis", 14, "Nader Khelifi", "EST-014", "2002-09-12", "Attaquant", "TUN", "Remplaçant", "", ""],
    ["ES Tunis", 18, "Bilel Sassi", "EST-018", "2001-04-04", "Milieu", "TUN", "Remplaçant", "", ""],
]

EVENTS_HEADER = ["Minute", "Equipe", "Joueur", "Evenement", "Carton", "Detail", "Entrant", "Sortant"]
EVENTS = [
    [18, "Club Africain", "Ali Mansour", "But", "", "frappe du droit", "", ""],
    [41, "ES Tunis", "Mahmoud Amri", "Carton jaune", "jaune", "faute tactique", "", ""],
    [52, "ES Tunis", "Youssef Gharbi", "But", "", "penalty", "", ""],
    [63, "Club Africain", "Hamza Maaloul", "Carton jaune", "jaune", "contestation", "", ""],
    [70, "Club Africain", "", "Remplacement", "", "changement offensif", "Soufiane Jaziri", "Firas Jouini"],
    [83, "ES Tunis", "Karim Chaabane", "Carton rouge", "rouge", "jeu dangereux", "", ""],
    [88, "Club Africain", "Mohamed Salem", "But", "", "tête sur corner", "", ""],
]


def make_docx():
    doc = Document()
    doc.add_heading("Feuille de match - Exemple", level=1)
    for line in LINES:
        doc.add_paragraph(line)
    doc.add_heading("Joueurs", level=2)
    table = doc.add_table(rows=1, cols=len(PLAYERS_HEADER))
    for i, h in enumerate(PLAYERS_HEADER):
        table.rows[0].cells[i].text = h
    for row in PLAYERS:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_heading("Événements", level=2)
    table = doc.add_table(rows=1, cols=len(EVENTS_HEADER))
    for i, h in enumerate(EVENTS_HEADER):
        table.rows[0].cells[i].text = h
    for row in EVENTS:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.save(BASE / "sample_match_sheet.docx")


def make_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Raw"
    ws.append(["Ligne"])
    for line in LINES:
        ws.append([line])
    ws_players = wb.create_sheet("Players")
    ws_players.append(PLAYERS_HEADER)
    for row in PLAYERS:
        ws_players.append(row)
    ws_events = wb.create_sheet("Events")
    ws_events.append(EVENTS_HEADER)
    for row in EVENTS:
        ws_events.append(row)
    wb.save(BASE / "sample_match_sheet.xlsx")


def make_pdf():
    c = canvas.Canvas(str(BASE / "sample_match_sheet.pdf"), pagesize=A4)
    width, height = A4
    y = height - 1.4 * cm
    c.setFont("Helvetica-Bold", 14)
    c.drawString(1.5 * cm, y, "Feuille de match - Exemple")
    y -= 0.7 * cm
    c.setFont("Helvetica", 8.5)
    for line in LINES:
        c.drawString(1.5 * cm, y, line)
        y -= 0.38 * cm
    y -= 0.2 * cm
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(1.5 * cm, y, " | ".join(PLAYERS_HEADER))
    y -= 0.35 * cm
    c.setFont("Helvetica", 7.4)
    for row in PLAYERS:
        if y < 2 * cm:
            c.showPage(); y = height - 1.5 * cm; c.setFont("Helvetica", 7.4)
        c.drawString(1.5 * cm, y, " | ".join(map(str, row)))
        y -= 0.30 * cm
    y -= 0.3 * cm
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(1.5 * cm, y, " | ".join(EVENTS_HEADER))
    y -= 0.35 * cm
    c.setFont("Helvetica", 7.6)
    for row in EVENTS:
        if y < 2 * cm:
            c.showPage(); y = height - 1.5 * cm; c.setFont("Helvetica", 7.6)
        c.drawString(1.5 * cm, y, " | ".join(map(str, row)))
        y -= 0.32 * cm
    c.save()


if __name__ == "__main__":
    make_docx()
    make_xlsx()
    make_pdf()
