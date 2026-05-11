from __future__ import annotations

import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document


@dataclass
class ExtractedContent:
    text: str
    file_type: str
    tables: list[dict[str, Any]] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)
    confidence: float = 0.0


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _confidence(text: str, tables_count: int = 0) -> float:
    length = len(text.strip())
    if length < 20:
        return 0.10
    if length < 400:
        return 0.45 + min(tables_count, 3) * 0.05
    if length < 1200:
        return 0.70 + min(tables_count, 3) * 0.05
    return min(0.95, 0.82 + min(tables_count, 5) * 0.02)


def extract_pdf(path: Path) -> ExtractedContent:
    warnings: list[str] = []
    text_parts: list[str] = []
    tables: list[dict[str, Any]] = []

    try:
        import fitz  # PyMuPDF

        with fitz.open(path) as doc:
            for page_index, page in enumerate(doc, start=1):
                page_text = page.get_text("text") or ""
                if page_text.strip():
                    text_parts.append(f"\n--- PAGE {page_index} ---\n{page_text}")
    except Exception as exc:  # pragma: no cover - best effort fallback
        warnings.append(f"PyMuPDF n'a pas pu lire le PDF: {exc}")

    try:
        import pdfplumber

        with pdfplumber.open(path) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                for table_index, table in enumerate(page.extract_tables() or [], start=1):
                    if not table:
                        continue
                    rows = [[cell or "" for cell in row] for row in table]
                    tables.append({"page": page_index, "table": table_index, "rows": rows})
                    text_parts.append(f"\n--- TABLE PDF page {page_index}.{table_index} ---")
                    text_parts.extend([" | ".join(row) for row in rows])
    except Exception as exc:  # pragma: no cover - table extraction is optional
        warnings.append(f"Extraction de tableaux PDF partielle: {exc}")

    text = "\n".join(text_parts).strip()
    if not text:
        warnings.append("Aucun texte lisible n'a été détecté. Le PDF est probablement scanné: ajouter OCR Tesseract/Google Vision/Azure pour la production.")
    return ExtractedContent(text=text, file_type="pdf", tables=tables, warnings=warnings, confidence=_confidence(text, len(tables)))


def extract_docx(path: Path) -> ExtractedContent:
    warnings: list[str] = []
    text_parts: list[str] = []
    tables: list[dict[str, Any]] = []
    doc = Document(path)
    for paragraph in doc.paragraphs:
        value = paragraph.text.strip()
        if value:
            text_parts.append(value)
    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append({"page": None, "table": table_index, "rows": rows})
            text_parts.append(f"\n--- TABLE WORD {table_index} ---")
            text_parts.extend([" | ".join(row) for row in rows])
    text = "\n".join(text_parts).strip()
    return ExtractedContent(text=text, file_type="docx", tables=tables, warnings=warnings, confidence=_confidence(text, len(tables)))


def extract_excel(path: Path) -> ExtractedContent:
    warnings: list[str] = []
    text_parts: list[str] = []
    tables: list[dict[str, Any]] = []
    try:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str, keep_default_na=False)
    except Exception as exc:
        return ExtractedContent(text="", file_type=path.suffix.lower().lstrip("."), warnings=[f"Excel non lu: {exc}"], confidence=0.0)
    for sheet_name, df in sheets.items():
        df = df.fillna("")
        rows = [list(map(str, df.columns.tolist()))] + df.astype(str).values.tolist()
        tables.append({"sheet": sheet_name, "rows": rows})
        text_parts.append(f"\n--- SHEET {sheet_name} ---")
        text_parts.extend([" | ".join(map(str, row)) for row in rows])
    text = "\n".join(text_parts).strip()
    return ExtractedContent(text=text, file_type="excel", tables=tables, warnings=warnings, confidence=_confidence(text, len(tables)))


def extract_csv(path: Path) -> ExtractedContent:
    warnings: list[str] = []
    try:
        df = pd.read_csv(path, dtype=str, keep_default_na=False)
        rows = [list(map(str, df.columns.tolist()))] + df.astype(str).values.tolist()
        text = "\n".join(" | ".join(map(str, row)) for row in rows)
        return ExtractedContent(text=text, file_type="csv", tables=[{"sheet": path.name, "rows": rows}], warnings=warnings, confidence=_confidence(text, 1))
    except Exception as exc:
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            text = ""
        warnings.append(f"CSV lu en mode texte brut: {exc}")
        return ExtractedContent(text=text, file_type="csv", warnings=warnings, confidence=_confidence(text))


def extract_document(path: Path) -> ExtractedContent:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".xlsx", ".xlsm", ".xls"}:
        return extract_excel(path)
    if suffix == ".csv":
        return extract_csv(path)
    if suffix == ".doc":
        return ExtractedContent(text="", file_type="doc", warnings=["Le format .doc ancien n'est pas supporté directement. Convertir en .docx ou PDF texte."], confidence=0.0)
    return ExtractedContent(text="", file_type=suffix.lstrip("."), warnings=["Format non supporté."], confidence=0.0)
